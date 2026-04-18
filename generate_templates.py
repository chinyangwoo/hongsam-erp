"""
generate_templates.py  ─  홍삼스파 ERP 문서 양식 파일 자동 생성
각 문서를 실제 Excel/DOCX/PDF 파일로 생성하여 docs/ 폴더에 저장
"""
import os, json, re, sys

# ── 경로 설정 ────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE, 'docs')
os.makedirs(DOCS_DIR, exist_ok=True)

# ── document_data.js 파싱 ────────────────────────────
with open(os.path.join(BASE, 'document_data.js'), 'r', encoding='utf-8') as f:
    content = f.read()

# DOC_TEMPLATES 배열 추출
match = re.search(r'const DOC_TEMPLATES\s*=\s*\[(.*?)\];', content, re.DOTALL)
if not match:
    print("ERROR: DOC_TEMPLATES not found"); sys.exit(1)

raw = match.group(1)
# 각 { ... } 객체 파싱
items = re.findall(r"\{[^}]+\}", raw)

templates = []
for item in items:
    d = {}
    for key in ['name', 'cat', 'date', 'user', 'size', 'ext']:
        m = re.search(rf"{key}:\s*'([^']*)'", item)
        if m: d[key] = m.group(1)
    if 'name' in d:
        templates.append(d)

print(f"총 {len(templates)}건 양식 생성 시작...")

# ── 엑셀 양식 생성 ───────────────────────────────────
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

HEADER_FILL = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
HEADER_FONT = Font(name='맑은 고딕', size=11, bold=True, color='FFFFFF')
TITLE_FONT = Font(name='맑은 고딕', size=16, bold=True, color='1F4E79')
NORMAL_FONT = Font(name='맑은 고딕', size=10)
THIN_BORDER = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)

# 카테고리별 엑셀 컬럼 헤더 매핑
EXCEL_HEADERS = {
    '지출결의서': ['No', '일자', '적요(내용)', '거래처', '금액(원)', '비고'],
    '수입결의서': ['No', '일자', '적요(내용)', '입금처', '금액(원)', '비고'],
    '법인카드': ['No', '사용일', '사용처', '사용자', '금액(원)', '결재구분', '비고'],
    '매출': ['No', '일자', '구분', '내용', '수량', '단가(원)', '금액(원)', '비고'],
    '예산': ['No', '항목', '1월', '2월', '3월', '4월', '5월', '6월', '7월', '8월', '9월', '10월', '11월', '12월', '합계'],
    '거래명세서': ['No', '품목명', '규격', '수량', '단가(원)', '공급가액', '세액', '합계'],
    '세금계산서': ['No', '발행일', '거래처명', '사업자번호', '공급가액', '세액', '합계', '비고'],
    '급여대장': ['No', '사번', '성명', '부서', '직급', '기본급', '식대', '교통비', '시간외수당', '지급총액', '소득세', '주민세', '국민연금', '건강보험', '고용보험', '공제총액', '실수령액'],
    '퇴직금': ['No', '사번', '성명', '입사일', '퇴직일', '재직일수', '3개월 평균임금', '퇴직금', '비고'],
    '근태기록': ['No', '사번', '성명', '부서', '근무일', '출근시간', '퇴근시간', '근무시간', '연장근무', '비고'],
    '연차': ['No', '사번', '성명', '부서', '발생일수', '사용일수', '잔여일수', '비고'],
    '인사기록': ['No', '사번', '성명', '부서', '직급', '입사일', '전화번호', '주소', '비고'],
    '점검': ['No', '점검항목', '점검기준', '점검결과', '양호', '불량', '조치사항', '점검자', '비고'],
    '재고': ['No', '품목코드', '품목명', '규격', '단위', '전일재고', '입고', '출고', '금일재고', '안전재고', '비고'],
    '발주서': ['No', '품목명', '규격', '단위', '수량', '단가(원)', '금액(원)', '납기일', '비고'],
    '관리대장': ['No', '일자', '구분', '내용', '수량', '담당자', '비고'],
    '체크리스트': ['No', '점검항목', '점검기준', 'O/X', '비고'],
    '보고서': ['No', '항목', '전월', '당월', '증감', '비고'],
    '조사표': ['No', '조사항목', '결과', '점수', '비고'],
    '정산서': ['No', '일자', '내용', '수량', '단가', '금액(원)', '비고'],
    '신청서': ['No', '신청일', '신청자', '부서', '내용', '사유', '승인', '비고'],
    '대장': ['No', '일자', '구분', '내용', '수량/금액', '담당자', '비고'],
    '일지': ['No', '일자', '시간', '점검항목', '점검결과', '조치사항', '점검자'],
    '평가표': ['No', '평가항목', '배점', '자기평가', '상사평가', '최종점수', '비고'],
    '계획서': ['No', '항목', '세부내용', '일정', '담당자', '예산', '비고'],
    '기본': ['No', '항목', '내용', '비고'],
}

def guess_header_type(name):
    """파일명으로 적절한 헤더 타입 추정"""
    name_lower = name.lower()
    for key in EXCEL_HEADERS:
        if key in name_lower:
            return key
    # 키워드 매핑
    kw_map = {
        '결의서': '지출결의서', '법인카드': '법인카드', '매출': '매출',
        '예산': '예산', '거래명세': '거래명세서', '세금계산': '세금계산서',
        '급여': '급여대장', '퇴직': '퇴직금', '근태': '근태기록',
        '연차': '연차', '인사': '인사기록', '점검': '점검', '재고': '재고',
        '발주': '발주서', '관리대장': '관리대장', '체크리스트': '체크리스트',
        '보고서': '보고서', '정산': '정산서', '신청': '신청서',
        '대장': '대장', '일지': '일지', '평가': '평가표', '계획': '계획서',
        '조사': '조사표',
    }
    for kw, htype in kw_map.items():
        if kw in name:
            return htype
    return '기본'


def create_excel(filepath, doc_name):
    """엑셀 양식 생성"""
    wb = Workbook()
    ws = wb.active
    ws.title = '양식'

    # 제목
    pretty_name = doc_name.replace('.xlsx', '').replace('.xls', '').replace('_', ' ')
    ws.merge_cells('A1:H1')
    cell = ws['A1']
    cell.value = f'홍삼스파 - {pretty_name}'
    cell.font = TITLE_FONT
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 40

    # 회사정보
    ws['A2'] = '회사명: (주)헤일로유니버스 / 홍삼스파'
    ws['A2'].font = Font(name='맑은 고딕', size=9, color='666666')
    ws['A3'] = f'작성일: ____년 ____월 ____일  |  작성자: ________________  |  부서: ________________'
    ws['A3'].font = Font(name='맑은 고딕', size=9, color='666666')
    ws.row_dimensions[4].height = 10  # 빈 줄

    # 헤더 추정
    htype = guess_header_type(pretty_name)
    headers = EXCEL_HEADERS.get(htype, EXCEL_HEADERS['기본'])

    # 헤더 행
    header_row = 5
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = THIN_BORDER

    # 빈 데이터 행 (20행)
    for r in range(header_row + 1, header_row + 21):
        for col in range(1, len(headers) + 1):
            cell = ws.cell(row=r, column=col, value='')
            cell.border = THIN_BORDER
            cell.font = NORMAL_FONT
            if col == 1:
                cell.value = r - header_row
                cell.alignment = Alignment(horizontal='center')

    # 열 너비 자동 조정
    for col, h in enumerate(headers, 1):
        ws.column_dimensions[chr(64 + min(col, 26))].width = max(len(h) * 2 + 4, 12)

    # 하단 서명란
    sign_row = header_row + 22
    ws.cell(row=sign_row, column=1, value='작성자: ____________ (인)').font = NORMAL_FONT
    ws.cell(row=sign_row, column=4, value='검토자: ____________ (인)').font = NORMAL_FONT
    ws.cell(row=sign_row, column=7, value='승인자: ____________ (인)').font = NORMAL_FONT

    wb.save(filepath)


# ── DOCX 양식 생성 ──────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(filepath, doc_name):
    """Word 문서 양식 생성"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    pretty_name = doc_name.replace('.docx', '').replace('.doc', '').replace('_', ' ')

    # 제목
    title = doc.add_heading(pretty_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 회사 정보
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(주)헤일로유니버스 / 홍삼스파')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    # 기본 정보 테이블
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    cells = [
        ('문서번호', 'HS-________-___'), ('작성일자', '____년 ____월 ____일'),
        ('작성부서', '________________'), ('작성자', '________________')
    ]
    for i, (label, value) in enumerate(cells):
        row = i // 2
        col = (i % 2) * 2
        table.cell(row, col).text = label
        table.cell(row, col).paragraphs[0].runs[0].bold = True if table.cell(row, col).paragraphs[0].runs else None
        table.cell(row, col + 1).text = value

    doc.add_paragraph('')

    # 본문 섹션들 (양식 종류에 따라)
    if '계약서' in pretty_name:
        sections = ['제1조 (목적)', '제2조 (계약기간)', '제3조 (계약금액)',
                     '제4조 (지급조건)', '제5조 (계약의 해지)', '제6조 (손해배상)',
                     '제7조 (비밀유지)', '제8조 (분쟁해결)', '제9조 (기타사항)']
        for s in sections:
            doc.add_heading(s, level=2)
            doc.add_paragraph('_' * 60)
            doc.add_paragraph('')
    elif '기획' in pretty_name or '계획' in pretty_name or '제안' in pretty_name:
        sections = ['1. 배경 및 목적', '2. 추진 내용', '3. 세부 실행 계획',
                     '4. 예상 효과', '5. 소요 예산', '6. 일정 계획', '7. 기타']
        for s in sections:
            doc.add_heading(s, level=2)
            doc.add_paragraph('_' * 60)
            doc.add_paragraph('')
    elif '보고서' in pretty_name:
        sections = ['1. 개요', '2. 현황 분석', '3. 문제점 및 개선사항',
                     '4. 추진 실적', '5. 향후 계획', '6. 첨부자료']
        for s in sections:
            doc.add_heading(s, level=2)
            doc.add_paragraph('_' * 60)
            doc.add_paragraph('')
    else:
        doc.add_heading('1. 제목', level=2)
        doc.add_paragraph('_' * 60)
        doc.add_paragraph('')
        doc.add_heading('2. 내용', level=2)
        doc.add_paragraph('_' * 60)
        doc.add_paragraph('')
        doc.add_heading('3. 비고', level=2)
        doc.add_paragraph('_' * 60)

    # 서명란
    doc.add_paragraph('')
    doc.add_paragraph('')
    sign_table = doc.add_table(rows=2, cols=3)
    sign_table.style = 'Table Grid'
    for i, role in enumerate(['작성자', '검토자', '승인자']):
        sign_table.cell(0, i).text = role
        sign_table.cell(1, i).text = '\n\n         (인)\n'

    doc.save(filepath)


# ── PDF 양식 생성 ──────────────────────────────────
from fpdf import FPDF

# 한글 폰트 경로
FONT_PATH = None
FONT_CANDIDATES = [
    r'C:\Windows\Fonts\malgun.ttf',
    r'C:\Windows\Fonts\NanumGothic.ttf',
    r'C:\Windows\Fonts\gulim.ttc',
]
for fp in FONT_CANDIDATES:
    if os.path.exists(fp):
        FONT_PATH = fp
        break

def create_pdf(filepath, doc_name):
    """PDF 양식 생성"""
    pdf = FPDF()
    pdf.add_page()

    if FONT_PATH:
        pdf.add_font('Korean', '', FONT_PATH, uni=True)
        pdf.set_font('Korean', '', 18)
    else:
        pdf.set_font('Helvetica', 'B', 18)

    pretty_name = doc_name.replace('.pdf', '').replace('_', ' ')

    # 제목
    pdf.cell(0, 15, pretty_name, ln=True, align='C')
    pdf.ln(5)

    if FONT_PATH:
        pdf.set_font('Korean', '', 9)
    else:
        pdf.set_font('Helvetica', '', 9)

    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, '(주)헤일로유니버스 / 홍삼스파', ln=True, align='C')
    pdf.ln(5)

    # 문서정보
    if FONT_PATH:
        pdf.set_font('Korean', '', 10)
    else:
        pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(0, 0, 0)

    pdf.cell(30, 8, '문서번호:', border=1)
    pdf.cell(60, 8, 'HS-________-___', border=1)
    pdf.cell(30, 8, '작성일자:', border=1)
    pdf.cell(60, 8, '____년 ____월 ____일', border=1, ln=True)
    pdf.cell(30, 8, '작성부서:', border=1)
    pdf.cell(60, 8, '________________', border=1)
    pdf.cell(30, 8, '작성자:', border=1)
    pdf.cell(60, 8, '________________', border=1, ln=True)

    pdf.ln(10)

    # 본문
    sections = ['1. 개요', '2. 목적', '3. 내용', '4. 세부사항', '5. 비고']

    if '매뉴얼' in pretty_name or '지침' in pretty_name:
        sections = ['1. 적용범위', '2. 관련법규', '3. 용어정의', '4. 책임과 권한',
                     '5. 세부절차', '6. 주의사항', '7. 비상조치', '8. 기록관리']
    elif '성적서' in pretty_name or '검사' in pretty_name:
        sections = ['1. 검사기관', '2. 검사일자', '3. 검사항목', '4. 검사결과',
                     '5. 판정기준', '6. 종합판정', '7. 특이사항']
    elif '규정' in pretty_name or '규칙' in pretty_name:
        sections = ['제1장 총칙', '제2장 적용범위', '제3장 세부규정',
                     '제4장 벌칙', '제5장 부칙']

    for section in sections:
        if FONT_PATH:
            pdf.set_font('Korean', '', 12)
        else:
            pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 10, section, ln=True)

        if FONT_PATH:
            pdf.set_font('Korean', '', 10)
        else:
            pdf.set_font('Helvetica', '', 10)
        pdf.cell(0, 8, '_' * 80, ln=True)
        pdf.ln(5)

    # 서명란
    pdf.ln(15)
    y = pdf.get_y()
    for i, role in enumerate(['작성자', '검토자', '승인자']):
        x = 20 + i * 60
        pdf.set_xy(x, y)
        pdf.cell(50, 8, f'{role}: _________ (인)', border=1, align='C')

    pdf.output(filepath)


# ── HWP → DOCX로 대체 생성 ──────────────────────────
def create_hwp_as_docx(filepath, doc_name):
    """HWP는 프로그래밍으로 생성이 어려우므로 DOCX로 대체 생성"""
    docx_path = filepath  # 이미 .docx 확장자로 변경됨
    create_docx(docx_path, doc_name)


# ── 전체 생성 ─────────────────────────────────────
success = 0
errors = 0

for doc in templates:
    name = doc['name']
    ext = doc.get('ext', '')

    # HWP → DOCX로 변환
    if ext == 'hwp':
        actual_name = name.replace('.hwp', '.docx')
    else:
        actual_name = name

    filepath = os.path.join(DOCS_DIR, actual_name)

    try:
        if ext in ('xlsx', 'xls'):
            actual_path = filepath.replace('.xls', '.xlsx')  # .xls → .xlsx
            create_excel(actual_path, name)
        elif ext == 'docx' or ext == 'doc':
            create_docx(filepath, name)
        elif ext == 'pdf':
            create_pdf(filepath, name)
        elif ext == 'hwp':
            create_hwp_as_docx(filepath, name)
        else:
            create_docx(filepath, name)  # 기타 → docx

        success += 1
        if success % 20 == 0:
            print(f'  ... {success}건 완료')
    except Exception as e:
        errors += 1
        print(f'  ERROR: {name} → {e}')

print(f'\n완료! 성공: {success}건, 오류: {errors}건')
print(f'파일 위치: {DOCS_DIR}')
